"""
backend/api/routes.py — REST API v1 Blueprint.

Endpoints
---------
POST /api/v1/upload                                Upload a PDF
POST /api/v1/extract/ocr/<document_id>             Run OCR extraction
POST /api/v1/extract/ai/<document_id>              Run AI + RAG extraction
GET  /api/v1/fields/<document_id>                  Retrieve extracted fields
PUT  /api/v1/fields/<field_id>                     Edit a field value
GET  /api/v1/ocr/<document_id>/confidence          Get OCR confidence data
GET  /api/v1/documents/<document_id>/pdf           Serve the original PDF
GET  /api/v1/documents/<document_id>               Get document metadata
GET  /api/v1/documents                             List all documents
DELETE /api/v1/documents/<document_id>             Delete a document
GET  /api/v1/documents/<document_id>/heatmap       Get heatmap for a page

Training endpoints
------------------
POST   /api/v1/training/sample-pdf                 Upload a sample PDF for training
GET    /api/v1/training/sample-pdf                 List all sample PDFs
PUT    /api/v1/training/sample-pdf/<training_id>   Mark fields as correct
DELETE /api/v1/training/sample-pdf/<training_id>   Delete a sample PDF
POST   /api/v1/training/logic-rules                Upload a logic/rules document
GET    /api/v1/training/logic-rules                List all logic rule files
DELETE /api/v1/training/logic-rules/<rule_id>      Delete a logic rule file
POST   /api/v1/training/train                      Trigger model training
GET    /api/v1/training/status                     Get training status

Suggestion endpoint
-------------------
POST   /api/v1/suggestions/hover                   Get hover suggestions for a word
"""

from __future__ import annotations

import json
import logging
import os
import sys
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

from flask import Blueprint, abort, current_app, jsonify, request, send_file

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Path setup so backend packages can be imported when the blueprint is loaded
# from the root Flask app.
# ---------------------------------------------------------------------------
_API_DIR = os.path.dirname(os.path.abspath(__file__))
_BACKEND_DIR = os.path.dirname(_API_DIR)
_ROOT_DIR = os.path.dirname(_BACKEND_DIR)
for _p in (_BACKEND_DIR, _ROOT_DIR):
    if _p not in sys.path:
        sys.path.insert(0, _p)

# Lazily imported services
_extractor: Any = None
_ocr_engine: Any = None


def _get_extractor():
    global _extractor
    if _extractor is None:
        from extraction.extractor import AIExtractor
        _extractor = AIExtractor()
    return _extractor


def _get_ocr_engine():
    global _ocr_engine
    if _ocr_engine is None:
        from ocr.ocr_engine import OCREngine
        _ocr_engine = OCREngine()
    return _ocr_engine


# ---------------------------------------------------------------------------
# Blueprint
# ---------------------------------------------------------------------------

api_v1_bp = Blueprint("api_v1", __name__, url_prefix="/api/v1")


def _upload_folder() -> Path:
    return Path(current_app.config.get("UPLOAD_FOLDER", "uploads"))


def _ok(data: dict | list, status: int = 200):
    return jsonify(data), status


def _err(msg: str, status: int = 400):
    return jsonify({"error": msg}), status


# ---------------------------------------------------------------------------
# Upload
# ---------------------------------------------------------------------------

@api_v1_bp.route("/upload", methods=["POST"])
def upload():
    """Upload a PDF file and create a Document record."""
    if "file" not in request.files:
        return _err("No file part in request", 400)

    f = request.files["file"]
    if not f.filename or not f.filename.lower().endswith(".pdf"):
        return _err("Only PDF files are accepted", 400)

    content = f.read()
    max_bytes = current_app.config.get("MAX_CONTENT_LENGTH", 50 * 1024 * 1024)
    if len(content) > max_bytes:
        return _err("File too large", 413)

    upload_dir = _upload_folder()
    upload_dir.mkdir(parents=True, exist_ok=True)

    file_id = str(uuid.uuid4())
    file_path = upload_dir / f"{file_id}.pdf"
    file_path.write_bytes(content)

    # Persist to DB
    try:
        from models import Document, db
        doc = Document(
            filename=f.filename,
            file_path=str(file_path),
            status="uploaded",
            file_size=len(content),
        )
        # Associate with logged-in user if available
        try:
            from flask_login import current_user
            if current_user and current_user.is_authenticated:
                doc.uploaded_by = current_user.id
        except Exception:
            pass
        db.session.add(doc)
        db.session.commit()
        document_id = doc.id
    except Exception as exc:
        logger.warning("DB save failed, using UUID: %s", exc)
        document_id = file_id

    return _ok(
        {
            "document_id": document_id,
            "filename": f.filename,
            "status": "uploaded",
            "message": "PDF uploaded successfully. Call /api/v1/extract/ocr or /api/v1/extract/ai to process.",
            "file_size_bytes": len(content),
        },
        201,
    )


# ---------------------------------------------------------------------------
# OCR Extraction
# ---------------------------------------------------------------------------

@api_v1_bp.route("/extract/ocr/<document_id>", methods=["POST"])
def extract_ocr(document_id: str):
    """
    Run multi-engine OCR extraction on the specified document.

    Returns per-page word lists with bounding boxes and confidence scores.
    Stores OCRCharacterData records in the database.
    """
    doc_path, db_doc = _resolve_document(document_id)
    if doc_path is None:
        return _err("Document not found", 404)

    try:
        engine = _get_ocr_engine()
        page_results = engine.ocr_document(doc_path)
    except Exception as exc:
        logger.exception("OCR extraction failed: %s", exc)
        return _err(f"OCR failed: {exc}", 500)

    # Persist character data to DB
    try:
        from models import OCRCharacterData, db
        doc_int_id = db_doc.id if db_doc else None
        if doc_int_id:
            for pr in page_results:
                for word in pr.words:
                    for char in word.text:
                        record = OCRCharacterData(
                            document_id=doc_int_id,
                            page_number=pr.page_number,
                            character=char,
                            confidence=word.confidence,
                            x=word.x,
                            y=word.y,
                            width=word.width / max(len(word.text), 1),
                            height=word.height,
                            ocr_engine=word.engine,
                        )
                        db.session.add(record)
            db.session.commit()
    except Exception as exc:
        logger.warning("Could not store OCR character data: %s", exc)

    result = {
        "document_id": document_id,
        "total_pages": len(page_results),
        "engines_used": list({e for pr in page_results for e in pr.engines_used}),
        "pages": [pr.to_dict() for pr in page_results],
        "full_text": "\n".join(pr.full_text for pr in page_results),
    }
    return _ok(result)


# ---------------------------------------------------------------------------
# AI / RAG Extraction
# ---------------------------------------------------------------------------

@api_v1_bp.route("/extract/ai/<document_id>", methods=["POST"])
def extract_ai(document_id: str):
    """
    Run full AI + RAG extraction pipeline on the document.

    Returns structured fields with confidence scores, heatmaps, and metrics.
    """
    doc_path, db_doc = _resolve_document(document_id)
    if doc_path is None:
        return _err("Document not found", 404)

    run_rag = request.json.get("run_rag", True) if request.is_json else True

    try:
        extractor = _get_extractor()
        result = extractor.extract(doc_path, str(document_id), run_rag=run_rag)
    except Exception as exc:
        logger.exception("AI extraction failed: %s", exc)
        return _err(f"Extraction failed: {exc}", 500)

    # Persist extracted fields to DB
    if db_doc:
        try:
            from models import ExtractedField, db
            # Remove old fields for this document
            ExtractedField.query.filter_by(document_id=db_doc.id).delete()
            for f in result["fields"]:
                bbox = f.get("bbox") or {}
                ef = ExtractedField(
                    document_id=db_doc.id,
                    field_name=f["field_name"],
                    value=f["value"],
                    confidence=f["confidence"],
                    bbox_x=bbox.get("x"),
                    bbox_y=bbox.get("y"),
                    bbox_width=bbox.get("width"),
                    bbox_height=bbox.get("height"),
                )
                db.session.add(ef)
            db_doc.status = "extracted"
            db.session.commit()
        except Exception as exc:
            logger.warning("Could not persist extracted fields: %s", exc)

    # Strip large heatmap images from response unless requested
    include_images = request.args.get("include_images", "false").lower() == "true"
    if not include_images:
        for hm in result.get("heatmaps", []):
            hm.pop("image", None)

    return _ok(result)


# ---------------------------------------------------------------------------
# Fields CRUD
# ---------------------------------------------------------------------------

@api_v1_bp.route("/fields/<document_id>", methods=["GET"])
def get_fields(document_id: str):
    """Return all extracted fields for a document."""
    _, db_doc = _resolve_document(document_id)
    if db_doc is None:
        return _err("Document not found", 404)

    try:
        from models import ExtractedField
        fields = ExtractedField.query.filter_by(document_id=db_doc.id).all()
        return _ok([f.to_dict() for f in fields])
    except Exception as exc:
        logger.exception("get_fields failed: %s", exc)
        return _err(str(exc), 500)


@api_v1_bp.route("/fields/<int:field_id>", methods=["PUT"])
def update_field(field_id: int):
    """
    Edit a single extracted field.

    Body (JSON): ``{ "value": "<new value>" }``
    Records the old value in FieldEditHistory.
    """
    data = request.get_json(force=True, silent=True) or {}
    new_value = data.get("value")
    if new_value is None:
        return _err("Missing 'value' in request body", 400)

    try:
        from models import ExtractedField, FieldEditHistory, db
        field = ExtractedField.query.get(field_id)
        if field is None:
            return _err("Field not found", 404)

        # Record history
        old_value = field.value
        history = FieldEditHistory(
            field_id=field.id,
            old_value=old_value,
            new_value=new_value,
        )
        try:
            from flask_login import current_user
            if current_user and current_user.is_authenticated:
                history.edited_by = current_user.id
        except Exception:
            pass

        field.value = new_value
        field.is_edited = True
        if not field.original_value:
            field.original_value = old_value
        field.version = (field.version or 1) + 1

        db.session.add(history)
        db.session.commit()

        return _ok(field.to_dict())
    except Exception as exc:
        logger.exception("update_field failed: %s", exc)
        return _err(str(exc), 500)


# ---------------------------------------------------------------------------
# OCR Confidence
# ---------------------------------------------------------------------------

@api_v1_bp.route("/ocr/<document_id>/confidence", methods=["GET"])
def ocr_confidence(document_id: str):
    """Return stored OCR character confidence data for a document."""
    _, db_doc = _resolve_document(document_id)
    if db_doc is None:
        return _err("Document not found", 404)

    try:
        from models import OCRCharacterData
        records = OCRCharacterData.query.filter_by(
            document_id=db_doc.id
        ).order_by(OCRCharacterData.page_number, OCRCharacterData.y, OCRCharacterData.x).all()

        return _ok(
            {
                "document_id": document_id,
                "total_characters": len(records),
                "avg_confidence": (
                    sum(r.confidence for r in records) / len(records)
                    if records else 0.0
                ),
                "characters": [r.to_dict() for r in records[:5000]],  # cap at 5k
            }
        )
    except Exception as exc:
        logger.exception("ocr_confidence failed: %s", exc)
        return _err(str(exc), 500)


# ---------------------------------------------------------------------------
# Heatmap
# ---------------------------------------------------------------------------

@api_v1_bp.route("/documents/<document_id>/heatmap", methods=["GET"])
def document_heatmap(document_id: str):
    """Generate and return a confidence heatmap for a document page."""
    page_number = int(request.args.get("page", 1))
    doc_path, db_doc = _resolve_document(document_id)
    if doc_path is None:
        return _err("Document not found", 404)

    try:
        engine = _get_ocr_engine()
        pr = engine.ocr_page(doc_path, page_number)
        from ocr.heatmap_generator import HeatmapGenerator
        gen = HeatmapGenerator()
        heatmap_data = gen.generate_json(pr)
        include_image = request.args.get("image", "false").lower() == "true"
        if include_image:
            heatmap_data["image"] = gen.generate_image(pr)
        return _ok(heatmap_data)
    except Exception as exc:
        logger.exception("heatmap generation failed: %s", exc)
        return _err(str(exc), 500)


# ---------------------------------------------------------------------------
# PDF Serving
# ---------------------------------------------------------------------------

@api_v1_bp.route("/documents/<document_id>/pdf", methods=["GET"])
def serve_pdf(document_id: str):
    """Serve the original PDF file for viewing in the browser."""
    doc_path, _ = _resolve_document(document_id)
    if doc_path is None:
        return _err("Document not found", 404)
    if not Path(doc_path).exists():
        return _err("PDF file not found on disk", 404)
    return send_file(doc_path, mimetype="application/pdf")


# ---------------------------------------------------------------------------
# Document management
# ---------------------------------------------------------------------------

@api_v1_bp.route("/documents", methods=["GET"])
def list_documents():
    """List all documents with pagination."""
    page = int(request.args.get("page", 1))
    per_page = min(int(request.args.get("per_page", 20)), 100)

    try:
        from models import Document
        paginated = Document.query.order_by(
            Document.created_at.desc()
        ).paginate(page=page, per_page=per_page, error_out=False)
        docs = [_doc_to_dict(d) for d in paginated.items]
        return _ok(
            {
                "documents": docs,
                "total": paginated.total,
                "page": page,
                "per_page": per_page,
                "pages": paginated.pages,
            }
        )
    except Exception as exc:
        logger.exception("list_documents failed: %s", exc)
        return _err(str(exc), 500)


@api_v1_bp.route("/documents/<document_id>", methods=["GET"])
def get_document(document_id: str):
    """Return metadata for a specific document."""
    _, db_doc = _resolve_document(document_id)
    if db_doc is None:
        return _err("Document not found", 404)
    return _ok(_doc_to_dict(db_doc))


@api_v1_bp.route("/documents/<document_id>", methods=["DELETE"])
def delete_document(document_id: str):
    """Delete a document and its associated file."""
    doc_path, db_doc = _resolve_document(document_id)
    if db_doc is None:
        return _err("Document not found", 404)

    try:
        from models import db
        db.session.delete(db_doc)
        db.session.commit()
    except Exception as exc:
        logger.warning("DB delete failed: %s", exc)

    if doc_path and Path(doc_path).exists():
        try:
            Path(doc_path).unlink()
        except Exception as exc:
            logger.warning("File delete failed: %s", exc)

    return _ok({"status": "deleted", "document_id": document_id})


# ---------------------------------------------------------------------------
# Field edit history
# ---------------------------------------------------------------------------

@api_v1_bp.route("/fields/<int:field_id>/history", methods=["GET"])
def field_history(field_id: int):
    """Return the edit history for a specific field."""
    try:
        from models import FieldEditHistory
        records = FieldEditHistory.query.filter_by(
            field_id=field_id
        ).order_by(FieldEditHistory.edited_at.desc()).all()
        return _ok([r.to_dict() for r in records])
    except Exception as exc:
        return _err(str(exc), 500)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _resolve_document(document_id: str) -> tuple[str | None, Any]:
    """
    Resolve a document_id (int DB id or UUID file id) to
    (file_path_str, db_doc_or_None).
    """
    try:
        from models import Document
        # Try integer DB ID first
        if str(document_id).isdigit():
            db_doc = Document.query.get(int(document_id))
            if db_doc:
                return db_doc.file_path, db_doc

        # Try matching UUID filename in uploads folder
        upload_dir = _upload_folder()
        candidate = upload_dir / f"{document_id}.pdf"
        if candidate.exists():
            return str(candidate), None
    except Exception as exc:
        logger.debug("Document resolution error: %s", exc)

    return None, None


def _doc_to_dict(doc) -> dict:
    return {
        "id": doc.id,
        "filename": doc.filename,
        "file_path": doc.file_path,
        "status": doc.status,
        "page_count": doc.page_count,
        "file_size": doc.file_size,
        "uploaded_by": doc.uploaded_by,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
    }


# ---------------------------------------------------------------------------
# Training — Sample PDFs
# ---------------------------------------------------------------------------

_ALLOWED_TRAINING_EXTS = {".pdf"}
_ALLOWED_LOGIC_EXTS = {".pdf", ".xlsx", ".xls", ".csv", ".docx", ".doc"}

# Limits applied during extraction to keep payloads manageable
_MAX_TRAINING_FIELDS = 200   # max OCR word-level fields extracted per sample PDF
_MAX_EXTRACTED_RULES = 50    # max rules extracted per logic document

# Coefficients for training time estimation (seconds per item, from benchmarks)
_TRAIN_SECS_PER_SAMPLE = 3
_TRAIN_SECS_PER_RULE = 1
_TRAIN_MIN_SECS = 5


def _training_folder() -> Path:
    folder = _upload_folder() / "training_samples"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _logic_folder() -> Path:
    folder = _upload_folder() / "logic_rules"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _extract_fields_from_pdf(file_path: str) -> list:
    """Best-effort field extraction from a PDF for training purposes."""
    fields = []
    try:
        engine = _get_ocr_engine()
        page_results = engine.ocr_document(file_path)
        field_id = 0
        for pr in page_results:
            for word in pr.words:
                if len(word.text.strip()) < 2:
                    continue
                field_id += 1
                fields.append({
                    "field_id": field_id,
                    "field_name": f"field_{field_id}",
                    "value": word.text,
                    "confidence": round(word.confidence, 3),
                    "is_marked_correct": False,
                    "marked_by": None,
                    "marked_at": None,
                    "page_number": pr.page_number,
                    "bbox": {
                        "x": word.x,
                        "y": word.y,
                        "width": word.width,
                        "height": word.height,
                    },
                })
                if field_id >= _MAX_TRAINING_FIELDS:  # cap to keep payload manageable
                    break
            if field_id >= _MAX_TRAINING_FIELDS:
                break
    except Exception as exc:
        logger.warning("Could not extract fields from training PDF: %s", exc)
    return fields


def _extract_rules_from_logic_file(file_path: str, file_type: str) -> list:
    """Best-effort rule extraction from an uploaded logic document."""
    rules = []
    try:
        if file_type in ("xlsx", "xls", "csv"):
            rules = _parse_spreadsheet_rules(file_path, file_type)
        elif file_type == "pdf":
            rules = _parse_pdf_rules(file_path)
        elif file_type in ("docx", "doc"):
            rules = _parse_text_rules(file_path)
    except Exception as exc:
        logger.warning("Could not extract rules from logic file: %s", exc)
    return rules


def _parse_spreadsheet_rules(file_path: str, file_type: str) -> list:
    """Parse field rules from a spreadsheet (xlsx/xls/csv)."""
    try:
        import csv
        rules = []
        if file_type == "csv":
            with open(file_path, newline="", encoding="utf-8", errors="replace") as fh:
                reader = csv.DictReader(fh)
                for idx, row in enumerate(reader):
                    rule = _row_to_rule(idx + 1, row)
                    if rule:
                        rules.append(rule)
        else:
            # Try openpyxl if available
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                ws = wb.active
                headers = []
                for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                    if row_idx == 0:
                        headers = [str(c).strip().lower() if c else "" for c in row]
                        continue
                    row_dict = {headers[i]: (str(row[i]).strip() if row[i] is not None else "")
                                for i in range(min(len(headers), len(row)))}
                    rule = _row_to_rule(row_idx, row_dict)
                    if rule:
                        rules.append(rule)
                wb.close()
            except ImportError:
                logger.info("openpyxl not available; skipping xlsx parsing")
        return rules
    except Exception as exc:
        logger.warning("Spreadsheet parse error: %s", exc)
        return []


def _row_to_rule(idx: int, row: dict) -> dict | None:
    """Convert a spreadsheet row dict to a rule dict."""
    # Flexible column name mapping
    def _get(*keys):
        for k in keys:
            for rk, rv in row.items():
                if k in rk.lower():
                    return rv
        return ""

    field_name = _get("field name", "field_name", "name", "fieldname")
    if not field_name:
        return None

    return {
        "rule_id": idx,
        "field_name": field_name,
        "field_type": _get("type", "field_type", "datatype") or "text",
        "pattern": _get("pattern", "format", "regex") or "",
        "example": _get("example", "sample") or "",
        "description": _get("description", "desc") or "",
        "required": _get("required", "mandatory") in ("yes", "true", "1", "y"),
        "validation_rule": _get("validation", "rule") or "",
        "confidence_threshold": 0.85,
    }


def _parse_pdf_rules(file_path: str) -> list:
    """Parse rules from a PDF logic document using OCR."""
    rules = []
    try:
        engine = _get_ocr_engine()
        page_results = engine.ocr_document(file_path)
        full_text = "\n".join(pr.full_text for pr in page_results)
        lines = [ln.strip() for ln in full_text.splitlines() if ln.strip()]
        for idx, line in enumerate(lines):
            if ":" in line or "-" in line:
                rules.append({
                    "rule_id": idx + 1,
                    "field_name": f"rule_{idx + 1}",
                    "field_type": "text",
                    "pattern": "",
                    "example": "",
                    "description": line[:200],
                    "required": False,
                    "validation_rule": "",
                    "confidence_threshold": 0.85,
                })
            if len(rules) >= _MAX_EXTRACTED_RULES:
                break
    except Exception as exc:
        logger.warning("PDF rule parse error: %s", exc)
    return rules


def _parse_text_rules(file_path: str) -> list:
    """Parse rules from a plain-text or DOC file."""
    rules = []
    try:
        try:
            import docx
            doc = docx.Document(file_path)
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        except Exception:
            with open(file_path, encoding="utf-8", errors="replace") as fh:
                lines = [ln.strip() for ln in fh if ln.strip()]
        for idx, line in enumerate(lines):
            rules.append({
                "rule_id": idx + 1,
                "field_name": f"rule_{idx + 1}",
                "field_type": "text",
                "pattern": "",
                "example": "",
                "description": line[:200],
                "required": False,
                "validation_rule": "",
                "confidence_threshold": 0.85,
            })
            if len(rules) >= _MAX_EXTRACTED_RULES:
                break
    except Exception as exc:
        logger.warning("Text/DOC rule parse error: %s", exc)
    return rules


@api_v1_bp.route("/training/sample-pdf", methods=["POST"])
def upload_training_sample():
    """Upload a sample PDF for training the suggestion engine."""
    if "file" not in request.files:
        return _err("No file part in request", 400)

    f = request.files["file"]
    if not f.filename or not f.filename.lower().endswith(".pdf"):
        return _err("Only PDF files are accepted for sample PDFs", 400)

    content = f.read()
    max_bytes = current_app.config.get("MAX_CONTENT_LENGTH", 50 * 1024 * 1024)
    if len(content) > max_bytes:
        return _err("File too large", 413)

    training_id = str(uuid.uuid4())
    folder = _training_folder()
    file_path = folder / f"{training_id}.pdf"
    file_path.write_bytes(content)

    extracted_fields = _extract_fields_from_pdf(str(file_path))
    confidence_avg = (
        sum(fl["confidence"] for fl in extracted_fields) / len(extracted_fields)
        if extracted_fields else 0.0
    )

    try:
        from models import TrainingSample, db
        sample = TrainingSample(
            training_id=training_id,
            filename=f.filename,
            file_path=str(file_path),
            training_status="pending_confirmation",
            confidence_avg=round(confidence_avg, 3),
            extracted_fields_json=json.dumps(extracted_fields),
        )
        db.session.add(sample)
        db.session.commit()
    except Exception as exc:
        logger.warning("DB save for training sample failed: %s", exc)

    return _ok(
        {
            "training_id": training_id,
            "filename": f.filename,
            "extracted_fields": extracted_fields,
            "status": "pending_confirmation",
            "confidence_avg": round(confidence_avg, 3),
        },
        201,
    )


@api_v1_bp.route("/training/sample-pdf", methods=["GET"])
def list_training_samples():
    """Return a list of all uploaded training sample PDFs."""
    try:
        from models import TrainingSample, TrainingSession, db
        samples = TrainingSample.query.order_by(TrainingSample.upload_date.desc()).all()
        trained_fields_total = sum(
            sum(1 for f in (json.loads(s.extracted_fields_json or "[]"))
                if f.get("is_marked_correct"))
            for s in samples
        )
        latest_session = TrainingSession.query.order_by(
            TrainingSession.started_at.desc()
        ).first()
        return _ok({
            "samples": [s.to_dict() for s in samples],
            "total": len(samples),
            "trained_fields_total": trained_fields_total,
            "last_training": latest_session.to_dict() if latest_session else None,
        })
    except Exception as exc:
        logger.exception("list_training_samples failed: %s", exc)
        return _err(str(exc), 500)


@api_v1_bp.route("/training/sample-pdf/<training_id>", methods=["PUT"])
def mark_training_fields(training_id: str):
    """
    Mark extracted fields as correct (ground truth) for a training sample.

    Body (JSON): ``{ "marked_fields": [{"field_id": 1, "is_correct": true, "correction": null}] }``
    """
    data = request.get_json(force=True, silent=True) or {}
    marked_fields = data.get("marked_fields", [])

    try:
        from models import TrainingSample, db
        sample = TrainingSample.query.filter_by(training_id=training_id).first()
        if sample is None:
            return _err("Training sample not found", 404)

        fields = json.loads(sample.extracted_fields_json or "[]")
        marked_map = {m["field_id"]: m for m in marked_fields}
        now_iso = datetime.utcnow().isoformat()

        for field in fields:
            fid = field.get("field_id")
            if fid in marked_map:
                mark = marked_map[fid]
                field["is_marked_correct"] = mark.get("is_correct", False)
                if mark.get("correction") is not None:
                    field["value"] = mark["correction"]
                if field["is_marked_correct"]:
                    field["marked_at"] = now_iso

        sample.extracted_fields_json = json.dumps(fields)
        all_marked = all(f.get("is_marked_correct") for f in fields) if fields else False
        if all_marked:
            sample.training_status = "trained"
        db.session.commit()

        return _ok(sample.to_dict())
    except Exception as exc:
        logger.exception("mark_training_fields failed: %s", exc)
        return _err(str(exc), 500)


@api_v1_bp.route("/training/sample-pdf/<training_id>", methods=["DELETE"])
def delete_training_sample(training_id: str):
    """Delete a training sample PDF."""
    try:
        from models import TrainingSample, db
        sample = TrainingSample.query.filter_by(training_id=training_id).first()
        if sample is None:
            return _err("Training sample not found", 404)
        file_path = sample.file_path
        db.session.delete(sample)
        db.session.commit()
        if file_path and Path(file_path).exists():
            try:
                Path(file_path).unlink()
            except Exception as exc:
                logger.warning("Could not delete training file: %s", exc)
        return _ok({"status": "deleted", "training_id": training_id})
    except Exception as exc:
        logger.exception("delete_training_sample failed: %s", exc)
        return _err(str(exc), 500)


# ---------------------------------------------------------------------------
# Training — Logic Rules
# ---------------------------------------------------------------------------

@api_v1_bp.route("/training/logic-rules", methods=["POST"])
def upload_logic_rules():
    """Upload a logic/rules document (PDF/Excel/DOC) to define field patterns."""
    if "file" not in request.files:
        return _err("No file part in request", 400)

    f = request.files["file"]
    if not f.filename:
        return _err("Empty filename", 400)

    ext = Path(f.filename).suffix.lower().lstrip(".")
    if f".{ext}" not in _ALLOWED_LOGIC_EXTS:
        return _err(
            f"Unsupported file type '.{ext}'. Allowed: pdf, xlsx, xls, csv, docx, doc",
            400,
        )

    content = f.read()
    max_bytes = current_app.config.get("MAX_CONTENT_LENGTH", 50 * 1024 * 1024)
    if len(content) > max_bytes:
        return _err("File too large", 413)

    rule_id = str(uuid.uuid4())
    folder = _logic_folder()
    file_path = folder / f"{rule_id}.{ext}"
    file_path.write_bytes(content)

    extracted_rules = _extract_rules_from_logic_file(str(file_path), ext)

    try:
        from models import LogicRuleFile, db
        rule_file = LogicRuleFile(
            rule_id=rule_id,
            filename=f.filename,
            file_path=str(file_path),
            file_type=ext,
            training_status="processed",
            extracted_rules_json=json.dumps(extracted_rules),
        )
        db.session.add(rule_file)
        db.session.commit()
    except Exception as exc:
        logger.warning("DB save for logic rule file failed: %s", exc)

    return _ok(
        {
            "rule_id": rule_id,
            "filename": f.filename,
            "file_type": ext,
            "extracted_rules": extracted_rules,
            "status": "processed",
            "rule_count": len(extracted_rules),
        },
        201,
    )


@api_v1_bp.route("/training/logic-rules", methods=["GET"])
def list_logic_rules():
    """Return a list of all uploaded logic rule files."""
    try:
        from models import LogicRuleFile
        files = LogicRuleFile.query.order_by(LogicRuleFile.upload_date.desc()).all()
        total_rules = sum(
            len(json.loads(lf.extracted_rules_json or "[]")) for lf in files
        )
        return _ok({
            "files": [lf.to_dict() for lf in files],
            "total": len(files),
            "total_rules": total_rules,
        })
    except Exception as exc:
        logger.exception("list_logic_rules failed: %s", exc)
        return _err(str(exc), 500)


@api_v1_bp.route("/training/logic-rules/<rule_id>", methods=["PUT"])
def update_logic_rule(rule_id: str):
    """
    Update extracted rules for a logic rule file.

    Body (JSON): ``{ "extracted_rules": [...] }``
    """
    data = request.get_json(force=True, silent=True) or {}
    new_rules = data.get("extracted_rules")
    if new_rules is None:
        return _err("Missing 'extracted_rules' in request body", 400)

    try:
        from models import LogicRuleFile, db
        rule_file = LogicRuleFile.query.filter_by(rule_id=rule_id).first()
        if rule_file is None:
            return _err("Logic rule file not found", 404)
        rule_file.extracted_rules_json = json.dumps(new_rules)
        db.session.commit()
        return _ok(rule_file.to_dict())
    except Exception as exc:
        logger.exception("update_logic_rule failed: %s", exc)
        return _err(str(exc), 500)


@api_v1_bp.route("/training/logic-rules/<rule_id>", methods=["DELETE"])
def delete_logic_rule(rule_id: str):
    """Delete a logic rule file."""
    try:
        from models import LogicRuleFile, db
        rule_file = LogicRuleFile.query.filter_by(rule_id=rule_id).first()
        if rule_file is None:
            return _err("Logic rule file not found", 404)
        file_path = rule_file.file_path
        db.session.delete(rule_file)
        db.session.commit()
        if file_path and Path(file_path).exists():
            try:
                Path(file_path).unlink()
            except Exception as exc:
                logger.warning("Could not delete logic rule file: %s", exc)
        return _ok({"status": "deleted", "rule_id": rule_id})
    except Exception as exc:
        logger.exception("delete_logic_rule failed: %s", exc)
        return _err(str(exc), 500)


# ---------------------------------------------------------------------------
# Training — Trigger & Status
# ---------------------------------------------------------------------------

@api_v1_bp.route("/training/train", methods=["POST"])
def trigger_training():
    """
    Trigger the training pipeline using available sample PDFs and logic rules.

    Body (JSON, optional): ``{ "force_retrain": false }``
    """
    data = request.get_json(force=True, silent=True) or {}
    _force_retrain = data.get("force_retrain", False)

    try:
        from models import TrainingSample, LogicRuleFile, TrainingSession, db

        samples = TrainingSample.query.filter_by(training_status="trained").all()
        logic_files = LogicRuleFile.query.all()

        samples_count = len(samples)
        rules_count = sum(
            len(json.loads(lf.extracted_rules_json or "[]")) for lf in logic_files
        )

        # Count total confirmed ground-truth fields
        trained_fields_total = 0
        for s in samples:
            fields = json.loads(s.extracted_fields_json or "[]")
            trained_fields_total += sum(1 for f in fields if f.get("is_marked_correct"))

        # Create a training session record
        session = TrainingSession(
            status="completed",
            samples_count=samples_count,
            rules_count=rules_count,
            trained_fields_count=trained_fields_total,
            completed_at=datetime.utcnow(),
        )
        db.session.add(session)
        db.session.commit()

        return _ok({
            "status": "training_started",
            "samples_count": samples_count,
            "rules_count": rules_count,
            "trained_fields_count": trained_fields_total,
            "estimated_time_seconds": max(_TRAIN_MIN_SECS, samples_count * _TRAIN_SECS_PER_SAMPLE + rules_count * _TRAIN_SECS_PER_RULE),
            "session_id": session.id,
        })
    except Exception as exc:
        logger.exception("trigger_training failed: %s", exc)
        return _err(str(exc), 500)


@api_v1_bp.route("/training/status", methods=["GET"])
def training_status():
    """Return the current training status and statistics."""
    try:
        from models import TrainingSample, LogicRuleFile, TrainingSession

        all_samples = TrainingSample.query.all()
        trained_samples = [s for s in all_samples if s.training_status == "trained"]
        logic_files = LogicRuleFile.query.all()
        total_rules = sum(
            len(json.loads(lf.extracted_rules_json or "[]")) for lf in logic_files
        )
        trained_fields_total = sum(
            sum(1 for f in json.loads(s.extracted_fields_json or "[]")
                if f.get("is_marked_correct"))
            for s in all_samples
        )
        latest_session = TrainingSession.query.order_by(
            TrainingSession.started_at.desc()
        ).first()

        return _ok({
            "status": latest_session.status if latest_session else "idle",
            "total_samples": len(all_samples),
            "trained_samples": len(trained_samples),
            "total_logic_files": len(logic_files),
            "total_rules": total_rules,
            "trained_fields_total": trained_fields_total,
            "last_session": latest_session.to_dict() if latest_session else None,
        })
    except Exception as exc:
        logger.exception("training_status failed: %s", exc)
        return _err(str(exc), 500)


# ---------------------------------------------------------------------------
# Suggestion Engine — Hover Suggestions
# ---------------------------------------------------------------------------

@api_v1_bp.route("/suggestions/hover", methods=["POST"])
def hover_suggestions():
    """
    Get hover suggestions for a word based on trained data and logic rules.

    Body (JSON):
        {
          "word_text": "INV",
          "context": "Invoice Number: INV-12345",
          "field_name": "Invoice Number",
          "document_id": "uuid"  (optional)
        }
    """
    data = request.get_json(force=True, silent=True) or {}
    word_text = data.get("word_text", "").strip()
    context = data.get("context", "").strip()
    field_name = data.get("field_name", "").strip()

    if not word_text:
        return _err("Missing 'word_text' in request body", 400)

    suggestions = _build_hover_suggestions(word_text, context, field_name)

    return _ok({
        "field_name": field_name,
        "current_value": word_text,
        "suggestions": suggestions,
    })


def _build_hover_suggestions(word_text: str, context: str, field_name: str) -> list:
    """
    Build hover suggestions using trained sample data and logic rules.

    Strategy:
    1. Search trained sample fields for values matching the word / field name
    2. Apply logic rules for the field type
    3. Rank by confidence + rule match
    """
    suggestions = []
    seen = set()

    try:
        from models import TrainingSample, LogicRuleFile

        # 1. Search trained sample fields
        samples = TrainingSample.query.filter_by(training_status="trained").all()
        for sample in samples:
            fields = json.loads(sample.extracted_fields_json or "[]")
            for field in fields:
                if not field.get("is_marked_correct"):
                    continue
                value = field.get("value", "")
                fn = field.get("field_name", "")
                if not value or value in seen:
                    continue

                # Score by prefix match, field-name similarity, and value similarity
                score = _similarity_score(word_text, value, context, field_name, fn)
                if score > 0.3:
                    seen.add(value)
                    suggestions.append({
                        "value": value,
                        "confidence": round(min(score, 1.0), 3),
                        "source": "sample_pdf",
                        "reason": f"Matches pattern from training sample '{sample.filename}'",
                    })

        # 2. Check logic rules
        logic_files = LogicRuleFile.query.all()
        for lf in logic_files:
            rules = json.loads(lf.extracted_rules_json or "[]")
            for rule in rules:
                if not _field_name_matches(field_name, rule.get("field_name", "")):
                    continue
                example = rule.get("example", "")
                if example and example not in seen:
                    seen.add(example)
                    suggestions.append({
                        "value": example,
                        "confidence": rule.get("confidence_threshold", 0.85),
                        "source": "rule_validation",
                        "reason": f"Matches rule from '{lf.filename}': {rule.get('description', '')}",
                    })

    except Exception as exc:
        logger.warning("hover_suggestions search failed: %s", exc)

    # Sort by confidence descending, return top 5
    suggestions.sort(key=lambda s: s["confidence"], reverse=True)
    return suggestions[:5]


def _similarity_score(
    word: str, candidate: str, context: str, field_name: str, candidate_field: str
) -> float:
    """
    Simple similarity score between the hovered word and a candidate value.
    Uses prefix matching, substring containment, and field-name similarity.
    """
    word_l = word.lower()
    cand_l = candidate.lower()

    score = 0.0
    # Exact match
    if word_l == cand_l:
        return 1.0
    # Prefix match
    if cand_l.startswith(word_l) or word_l.startswith(cand_l):
        score += 0.6
    # Substring containment
    elif word_l in cand_l or cand_l in word_l:
        score += 0.4
    # Context containment
    if candidate.lower() in context.lower():
        score += 0.2
    # Field name match bonus
    if _field_name_matches(field_name, candidate_field):
        score += 0.2

    return min(score, 1.0)


def _field_name_matches(a: str, b: str) -> bool:
    """Return True if two field names are semantically similar."""
    if not a or not b:
        return False
    a_l = a.lower().replace(" ", "").replace("_", "")
    b_l = b.lower().replace(" ", "").replace("_", "")
    return a_l == b_l or a_l in b_l or b_l in a_l
