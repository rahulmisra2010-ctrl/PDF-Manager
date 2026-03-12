"""
backend/services/extraction_pipeline.py — Multi-tool extraction pipeline.

Orchestrates all available extraction tools in an intelligent order with
confidence scoring and fallback chains for maximum extraction success.

Tool chain order
----------------
1. Mindee IDP        — best for structured documents (invoices, receipts)
2. Koncile.ai        — enterprise IDP for complex documents
3. PyMuPDF           — direct PDF text + AcroForm widget extraction
4. pdfplumber        — layout-aware PDF text extraction
5. Tesseract OCR     — scanned / image-only documents
6. python-docx       — Word (.docx) documents
7. openpyxl          — Excel (.xlsx / .xls) files
8. LLM (GPT)         — semantic field extraction from OCR text
9. Regex patterns    — fallback for known field formats
"""

from __future__ import annotations

import io
import json
import logging
import os
import re
import tempfile
from typing import Any

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Regex fallback patterns for common field types
# ---------------------------------------------------------------------------
_REGEX_PATTERNS: dict[str, str] = {
    "Email": r"[\w.+-]+@[\w-]+\.[a-z]{2,}",
    "Phone": r"(?:\+\d{1,3}[\s-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}",
    "Date": r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
    "Amount": r"\$\s?\d{1,3}(?:,\d{3})*(?:\.\d{2})?",
    "ZIP Code": r"\b\d{5}(?:-\d{4})?\b",
    "Invoice Number": r"(?:INV|Invoice)\s*[#:\-]?\s*([A-Z0-9\-]+)",
}


def _apply_regex_fallback(text: str) -> dict[str, str]:
    """Extract common fields using regex patterns as a last resort."""
    results: dict[str, str] = {}
    for field_name, pattern in _REGEX_PATTERNS.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            results[field_name] = match.group(0).strip()
    return results


# ---------------------------------------------------------------------------
# Individual tool extractors
# ---------------------------------------------------------------------------

def extract_with_mindee(file_data: bytes, filename: str, api_key: str) -> dict[str, Any]:
    """
    Extract using Mindee IDP API.

    Returns ``{"fields": {...}, "confidence": {...}, "tool": "mindee", "raw": ""}``
    """
    try:
        from mindee import Client, product  # type: ignore[import]

        mindee_client = Client(api_key=api_key)
        input_doc = mindee_client.source_from_bytes(file_data, filename)

        # Use InvoiceV4 as the general-purpose model; fallback gracefully
        result = mindee_client.parse(product.InvoiceV4, input_doc)
        prediction = result.document.inference.prediction

        fields: dict[str, str] = {}
        confidence: dict[str, float] = {}

        field_map = {
            "invoice_number": "Invoice Number",
            "date": "Invoice Date",
            "due_date": "Due Date",
            "supplier_name": "Supplier",
            "supplier_address": "Supplier Address",
            "customer_name": "Customer",
            "customer_address": "Customer Address",
            "total_net": "Net Amount",
            "total_amount": "Total Amount",
            "total_tax": "Tax Amount",
        }

        for attr, label in field_map.items():
            try:
                val = getattr(prediction, attr, None)
                if val is not None:
                    str_val = str(val.value) if hasattr(val, "value") else str(val)
                    conf = float(val.confidence) if hasattr(val, "confidence") else 0.8
                    if str_val and str_val.lower() not in ("none", ""):
                        fields[label] = str_val
                        confidence[label] = conf
            except Exception:
                pass

        logger.info("Mindee extracted %d fields", len(fields))
        return {"fields": fields, "confidence": confidence, "tool": "mindee", "raw": ""}

    except ImportError:
        logger.debug("mindee package not installed; skipping")
        return {}
    except Exception as exc:
        logger.warning("Mindee extraction failed: %s", exc)
        return {}


def extract_with_koncile(file_data: bytes, filename: str, api_key: str) -> dict[str, Any]:
    """
    Extract using Koncile.ai IDP platform.

    Returns ``{"fields": {...}, "confidence": {...}, "tool": "koncile", "raw": ""}``
    """
    try:
        import requests  # type: ignore[import]

        # Koncile REST API upload endpoint (adjust base URL if needed)
        base_url = os.environ.get("KONCILE_BASE_URL", "https://api.koncile.ai")
        headers = {"Authorization": f"Bearer {api_key}"}

        upload_resp = requests.post(
            f"{base_url}/v1/documents",
            headers=headers,
            files={"file": (filename, file_data)},
            timeout=60,
        )
        upload_resp.raise_for_status()
        doc_id = upload_resp.json().get("document_id") or upload_resp.json().get("id")

        # Poll for extraction result
        import time

        for _ in range(20):
            status_resp = requests.get(
                f"{base_url}/v1/documents/{doc_id}",
                headers=headers,
                timeout=30,
            )
            status_resp.raise_for_status()
            data = status_resp.json()
            if data.get("status") in ("completed", "done"):
                break
            time.sleep(3)

        raw_fields = data.get("fields") or data.get("extracted_fields") or {}
        fields: dict[str, str] = {}
        confidence: dict[str, float] = {}

        for key, val in raw_fields.items():
            if isinstance(val, dict):
                fields[key] = str(val.get("value", ""))
                confidence[key] = float(val.get("confidence", 0.7))
            else:
                fields[key] = str(val)
                confidence[key] = 0.7

        logger.info("Koncile extracted %d fields", len(fields))
        return {"fields": fields, "confidence": confidence, "tool": "koncile", "raw": ""}

    except ImportError:
        logger.debug("requests not available; skipping Koncile")
        return {}
    except Exception as exc:
        logger.warning("Koncile extraction failed: %s", exc)
        return {}


def extract_with_pymupdf(file_data: bytes) -> dict[str, Any]:
    """
    Extract text and AcroForm widget fields using PyMuPDF.

    Returns ``{"fields": {...}, "confidence": {...}, "tool": "pymupdf", "raw": "..."}``
    """
    try:
        import fitz  # type: ignore[import] — PyMuPDF

        doc = fitz.open(stream=file_data, filetype="pdf")

        # --- AcroForm widget fields ---
        widget_fields: dict[str, str] = {}
        for page in doc:
            for widget in page.widgets() or []:
                name = (widget.field_name or "").strip()
                value = str(widget.field_value or "").strip()
                if name and value:
                    widget_fields[name] = value

        if widget_fields:
            doc.close()
            logger.info("PyMuPDF AcroForm: %d widget fields", len(widget_fields))
            return {
                "fields": widget_fields,
                "confidence": {k: 0.95 for k in widget_fields},
                "tool": "pymupdf",
                "raw": "",
            }

        # --- Plain text extraction ---
        raw_text = "\n".join(page.get_text() for page in doc)
        doc.close()

        logger.info("PyMuPDF text: %d chars", len(raw_text))
        return {
            "fields": {},
            "confidence": {},
            "tool": "pymupdf",
            "raw": raw_text,
        }

    except ImportError:
        logger.debug("PyMuPDF not installed; skipping")
        return {}
    except Exception as exc:
        logger.warning("PyMuPDF extraction failed: %s", exc)
        return {}


def extract_with_pdfplumber(file_data: bytes) -> dict[str, Any]:
    """
    Extract text using pdfplumber (layout-aware).

    Returns ``{"fields": {}, "confidence": {}, "tool": "pdfplumber", "raw": "..."}``
    """
    try:
        import pdfplumber  # type: ignore[import]

        with pdfplumber.open(io.BytesIO(file_data)) as pdf:
            raw_text = "\n".join(
                page.extract_text() or "" for page in pdf.pages
            )

        logger.info("pdfplumber text: %d chars", len(raw_text))
        return {
            "fields": {},
            "confidence": {},
            "tool": "pdfplumber",
            "raw": raw_text,
        }

    except ImportError:
        logger.debug("pdfplumber not installed; skipping")
        return {}
    except Exception as exc:
        logger.warning("pdfplumber extraction failed: %s", exc)
        return {}


def extract_with_tesseract(file_data: bytes, filename: str) -> dict[str, Any]:
    """
    Extract text using Tesseract OCR.  Handles PDFs (via pdf2image) and images.

    Returns ``{"fields": {}, "confidence": {}, "tool": "tesseract", "raw": "..."}``
    """
    try:
        import pytesseract  # type: ignore[import]
        from PIL import Image  # type: ignore[import]

        ext = os.path.splitext(filename)[1].lower()
        images = []

        if ext == ".pdf":
            try:
                from pdf2image import convert_from_bytes  # type: ignore[import]

                images = convert_from_bytes(file_data)
            except ImportError:
                logger.debug("pdf2image not installed; cannot OCR PDF pages")
                return {}
        else:
            images = [Image.open(io.BytesIO(file_data))]

        raw_text = ""
        for img in images:
            raw_text += pytesseract.image_to_string(img) + "\n"

        logger.info("Tesseract OCR: %d chars", len(raw_text))
        return {
            "fields": {},
            "confidence": {},
            "tool": "tesseract",
            "raw": raw_text,
        }

    except ImportError:
        logger.debug("pytesseract not installed; skipping")
        return {}
    except Exception as exc:
        logger.warning("Tesseract extraction failed: %s", exc)
        return {}


def extract_with_docx(file_data: bytes) -> dict[str, Any]:
    """
    Extract text from Word (.docx) documents using python-docx.

    Returns ``{"fields": {}, "confidence": {}, "tool": "docx", "raw": "..."}``
    """
    try:
        from docx import Document  # type: ignore[import]

        doc = Document(io.BytesIO(file_data))
        raw_text = "\n".join(para.text for para in doc.paragraphs)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                raw_text += "\t".join(cell.text for cell in row.cells) + "\n"

        logger.info("python-docx text: %d chars", len(raw_text))
        return {
            "fields": {},
            "confidence": {},
            "tool": "docx",
            "raw": raw_text,
        }

    except ImportError:
        logger.debug("python-docx not installed; skipping")
        return {}
    except Exception as exc:
        logger.warning("python-docx extraction failed: %s", exc)
        return {}


def extract_with_openpyxl(file_data: bytes) -> dict[str, Any]:
    """
    Extract data from Excel (.xlsx) files using openpyxl.

    Returns field-value pairs derived from the first two columns.
    Returns ``{"fields": {...}, "confidence": {...}, "tool": "openpyxl", "raw": "..."}``
    """
    try:
        import openpyxl  # type: ignore[import]

        wb = openpyxl.load_workbook(io.BytesIO(file_data), data_only=True)
        fields: dict[str, str] = {}
        confidence: dict[str, float] = {}
        raw_lines: list[str] = []

        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                if not any(row):
                    continue
                row_text = "\t".join(str(c) if c is not None else "" for c in row)
                raw_lines.append(row_text)
                # Treat first two columns as key→value pairs
                if len(row) >= 2 and row[0] and row[1]:
                    key = str(row[0]).strip()
                    val = str(row[1]).strip()
                    if key:
                        fields[key] = val
                        confidence[key] = 0.85

        logger.info("openpyxl extracted %d fields", len(fields))
        return {
            "fields": fields,
            "confidence": confidence,
            "tool": "openpyxl",
            "raw": "\n".join(raw_lines),
        }

    except ImportError:
        logger.debug("openpyxl not installed; skipping")
        return {}
    except Exception as exc:
        logger.warning("openpyxl extraction failed: %s", exc)
        return {}


def extract_with_llm(raw_text: str, api_key: str) -> dict[str, Any]:
    """
    Extract structured fields using an LLM (OpenAI GPT).

    Returns ``{"fields": {...}, "confidence": {...}, "tool": "llm", "raw": ""}``
    """
    if not raw_text or not api_key:
        return {}
    try:
        from openai import OpenAI  # type: ignore[import]

        client = OpenAI(api_key=api_key)

        # Truncate to avoid token limits
        truncated = raw_text[:4000]
        model = os.environ.get("OPENAI_MODEL", "gpt-3.5-turbo")

        prompt = (
            "Extract all key-value pairs from the following document text.\n"
            "Return ONLY valid JSON (object) with field names as keys and "
            "string values. Do not include any commentary.\n\n"
            f"Document:\n{truncated}\n\nJSON:"
        )

        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=1000,
        )
        raw_json = response.choices[0].message.content or "{}"

        # Strip markdown fences if present
        raw_json = re.sub(r"```(?:json)?", "", raw_json).strip().strip("`")

        fields = json.loads(raw_json)
        if not isinstance(fields, dict):
            fields = {}

        confidence = {k: 0.75 for k in fields}
        logger.info("LLM extracted %d fields", len(fields))
        return {"fields": fields, "confidence": confidence, "tool": "llm", "raw": ""}

    except ImportError:
        logger.debug("openai package not installed; skipping LLM extraction")
        return {}
    except Exception as exc:
        logger.warning("LLM extraction failed: %s", exc)
        return {}


# ---------------------------------------------------------------------------
# Document type detection
# ---------------------------------------------------------------------------

_DOC_TYPE_KEYWORDS: dict[str, list[str]] = {
    "invoice": ["invoice", "inv #", "bill to", "amount due", "total due"],
    "receipt": ["receipt", "thank you for your purchase", "subtotal", "cashier"],
    "form": ["please fill", "signature", "date of birth", "applicant"],
    "report": ["report", "summary", "analysis", "findings", "conclusion"],
    "contract": ["agreement", "terms and conditions", "parties", "clause", "whereas"],
    "resume": ["experience", "education", "skills", "references", "objective"],
    "letter": ["dear", "sincerely", "regards", "to whom it may concern"],
}


def detect_document_type(text: str) -> str:
    """Return a human-readable document type label based on keyword matching."""
    lower = text.lower()
    scores: dict[str, int] = {}
    for doc_type, keywords in _DOC_TYPE_KEYWORDS.items():
        scores[doc_type] = sum(1 for kw in keywords if kw in lower)
    best = max(scores, key=lambda k: scores[k])
    return best if scores[best] > 0 else "unknown"


# ---------------------------------------------------------------------------
# Result merger
# ---------------------------------------------------------------------------

def _merge_results(results: list[dict[str, Any]]) -> dict[str, Any]:
    """
    Merge extraction results from multiple tools.

    For each field, keep the value with the highest confidence.  Overall
    quality score is the mean confidence of all accepted fields.
    """
    merged_fields: dict[str, str] = {}
    merged_confidence: dict[str, float] = {}
    best_tool: str = ""
    best_tool_count = 0

    for result in results:
        if not result:
            continue
        fields = result.get("fields") or {}
        confidence = result.get("confidence") or {}
        tool = result.get("tool", "unknown")

        for field, value in fields.items():
            conf = float(confidence.get(field, 0.5))
            if field not in merged_confidence or conf > merged_confidence[field]:
                merged_fields[field] = value
                merged_confidence[field] = conf

        if len(fields) > best_tool_count:
            best_tool_count = len(fields)
            best_tool = tool

    quality = (
        sum(merged_confidence.values()) / len(merged_confidence)
        if merged_confidence
        else 0.0
    )

    return {
        "fields": merged_fields,
        "confidence": merged_confidence,
        "tool": best_tool,
        "quality_score": round(quality, 4),
    }


# ---------------------------------------------------------------------------
# Public pipeline entry point
# ---------------------------------------------------------------------------

def run_extraction_pipeline(
    file_data: bytes,
    filename: str,
    *,
    mindee_key: str | None = None,
    koncile_key: str | None = None,
    openai_key: str | None = None,
    use_ocr: bool = True,
) -> dict[str, Any]:
    """
    Run all applicable extraction tools and return merged results.

    Parameters
    ----------
    file_data:
        Raw bytes of the uploaded file.
    filename:
        Original filename including extension (used to detect file type).
    mindee_key:
        Mindee API key — omit or pass ``None`` to skip.
    koncile_key:
        Koncile API key — omit or pass ``None`` to skip.
    openai_key:
        OpenAI API key for LLM-based extraction — omit to skip.
    use_ocr:
        If ``True``, Tesseract OCR is attempted when other tools yield no text.

    Returns
    -------
    dict with keys:
        ``fields``         — merged field dict  
        ``confidence``     — per-field confidence scores  
        ``tool``           — name of the dominant extraction tool  
        ``quality_score``  — average confidence (0–1)  
        ``document_type``  — detected document type  
        ``raw_text``       — best raw text extracted  
    """
    ext = os.path.splitext(filename)[1].lower()
    results: list[dict[str, Any]] = []
    raw_texts: list[str] = []

    # --- IDP cloud tools (highest confidence for structured docs) ---
    if mindee_key and ext in (".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".webp"):
        r = extract_with_mindee(file_data, filename, mindee_key)
        if r:
            results.append(r)

    if koncile_key:
        r = extract_with_koncile(file_data, filename, koncile_key)
        if r:
            results.append(r)

    # --- File-type specific tools ---
    if ext == ".pdf":
        r = extract_with_pymupdf(file_data)
        if r:
            results.append(r)
            if r.get("raw"):
                raw_texts.append(r["raw"])

        r = extract_with_pdfplumber(file_data)
        if r:
            results.append(r)
            if r.get("raw"):
                raw_texts.append(r["raw"])

    elif ext in (".docx",):
        r = extract_with_docx(file_data)
        if r:
            results.append(r)
            if r.get("raw"):
                raw_texts.append(r["raw"])

    elif ext in (".xlsx", ".xls"):
        r = extract_with_openpyxl(file_data)
        if r:
            results.append(r)
            if r.get("raw"):
                raw_texts.append(r["raw"])

    elif ext in (".txt",):
        # Plain text files — read directly
        try:
            raw_text = file_data.decode("utf-8", errors="replace")
            if raw_text.strip():
                raw_texts.append(raw_text)
                results.append({
                    "fields": {},
                    "confidence": {},
                    "tool": "plaintext",
                    "raw": raw_text,
                })
        except Exception as exc:
            logger.warning("Plain text read failed: %s", exc)

    # --- OCR for images and scanned PDFs ---
    if use_ocr and ext in (".pdf", ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"):
        r = extract_with_tesseract(file_data, filename)
        if r:
            results.append(r)
            if r.get("raw"):
                raw_texts.append(r["raw"])

    # --- LLM over raw text ---
    best_raw = max(raw_texts, key=len) if raw_texts else ""
    if best_raw and openai_key:
        r = extract_with_llm(best_raw, openai_key)
        if r:
            results.append(r)

    # --- Regex fallback ---
    if best_raw:
        regex_fields = _apply_regex_fallback(best_raw)
        if regex_fields:
            results.append({
                "fields": regex_fields,
                "confidence": {k: 0.5 for k in regex_fields},
                "tool": "regex",
                "raw": "",
            })

    # --- Merge all results ---
    merged = _merge_results(results)
    doc_type = detect_document_type(best_raw)

    return {
        "fields": merged.get("fields", {}),
        "confidence": merged.get("confidence", {}),
        "tool": merged.get("tool", "none"),
        "quality_score": merged.get("quality_score", 0.0),
        "document_type": doc_type,
        "raw_text": best_raw[:5000],  # cap raw text stored
    }
