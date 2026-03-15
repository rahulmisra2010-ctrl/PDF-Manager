"""
blueprints/training.py — Training examples API and UI blueprint.

Routes
------
POST   /api/v1/training/add              — save labeled training examples for a document
GET    /api/v1/training/examples         — list all training examples as JSON
GET    /api/v1/training/list             — alias for /api/v1/training/examples
DELETE /api/v1/training/<id>             — delete a single training example by ID
POST   /api/v1/training/save-roi         — save multiple ROI-annotated examples at once
POST   /api/v1/training/apply/<doc_id>   — overwrite all document fields with training data
GET    /training/examples                — HTML view of all training examples
GET    /training/upload-sample           — form UI to create a sample training entry
POST   /training/upload-sample           — process the uploaded sample form (file → review, manual → save)
GET    /training/upload-sample/review    — review extracted fields before saving
POST   /training/upload-sample/review    — confirm and save reviewed fields
POST   /training/examples/<id>/delete   — delete a training example group (document)
POST /api/v1/training/add         — save labeled training examples for a document
GET  /api/v1/training/examples    — list all training examples as JSON
GET  /training/examples           — HTML view of all training examples
GET  /training/upload-sample      — form UI to create a sample training entry
POST /training/upload-sample      — process the uploaded sample form
POST /training/extract-preview    — extract fields from a file for preview (JSON)
POST /training/examples/<id>/delete — delete a training example group (document)
"""

from __future__ import annotations

import io
import logging
import os
import re

from flask import Blueprint, current_app, flash, jsonify, redirect, render_template, request, session as flask_session, url_for
from flask_login import current_user, login_required
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from models import AuditLog, Document, ExtractedField, TrainingExample, db

logger = logging.getLogger(__name__)

# Document types available for the upload-sample form
DOCUMENT_TYPES = [
    "Address Book",
    "Invoice",
    "Receipt",
    "Contract",
    "Resume",
    "Medical Record",
    "Tax Form",
    "Bank Statement",
    "Other",
]

training_bp = Blueprint("training", __name__)

# Allowed file extensions for the file-upload mode
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx", ".xlsx", ".xls"}

# Well-known address-book / form field names used for inline-keyword matching.
# Ordered longest-first so that "Street Address" is matched before "Address".
_KNOWN_FIELD_NAMES: list[str] = [
    "Street Address",
    "Cell Phone",
    "Home Phone",
    "Work Phone",
    "Email Address",
    "First Name",
    "Last Name",
    "Full Name",
    "Middle Name",
    "Zip Code",
    "Postal Code",
    "Date of Birth",
    "Name",
    "Address",
    "Street",
    "City",
    "State",
    "Zip",
    "Country",
    "Phone",
    "Mobile",
    "Email",
    "Company",
    "Organization",
    "Title",
    "Fax",
    "Website",
    "Notes",
    "Birthday",
]


def _extract_fields_from_file(file_storage: FileStorage) -> dict[str, str]:
    """Extract field name → value pairs from an uploaded file.

    Supports PDF, TXT, DOCX, and XLSX/XLS files.
    For generic files the parser looks for lines matching ``Field: Value``.
    Returns an ordered dict (insertion-ordered in Python 3.7+).
    """
    filename = file_storage.filename or ""
    ext = os.path.splitext(filename)[1].lower()
    data = file_storage.read()

    if ext == ".txt":
        return _parse_txt(data.decode("utf-8", errors="replace"))

    if ext == ".pdf":
        return _parse_pdf(data)

    if ext == ".docx":
        return _parse_docx(data)

    if ext in (".xlsx", ".xls"):
        return _parse_excel(data, ext)

    return {}


def _parse_txt(text: str) -> dict[str, str]:
    """Parse ``Field: Value`` pairs from plain text (one per line)."""
    fields: dict[str, str] = {}
    for line in text.splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        # Accept lines like "Name: Rahul Misra" or "Name = Rahul Misra"
        m = re.match(r"^([^:=]+?)\s*[:=]\s*(.+)$", line)
        if m:
            fname = m.group(1).strip()
            fval = m.group(2).strip()
            if fname and fval:
                fields[fname] = fval
    return fields


def _extract_pdf_text(data: bytes) -> str:
    """Extract full text from PDF bytes, trying pdfplumber then PyMuPDF."""
    # Try pdfplumber first (better layout-aware extraction)
    try:
        import pdfplumber  # type: ignore[import]

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            parts: list[str] = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                parts.append(page_text)
        return "\n".join(parts)
    except Exception:
        pass

    # Fallback: PyMuPDF
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=data, filetype="pdf")
        parts = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("_extract_pdf_text: both extractors failed: %s", exc)
        return ""


def _parse_known_fields_inline(text: str) -> dict[str, str]:
    """Match known field names that appear on the same line as their value.

    Handles patterns like::

        Name Rahul Misra
        Cell Phone 7699888010
        Zip Code 713301

    The value is everything after the field name on the same line, stripped of
    leading underscores/dashes used as placeholder fill-lines.
    """
    fields: dict[str, str] = {}
    for field in _KNOWN_FIELD_NAMES:
        if field in fields:
            continue  # already captured via a longer match
        pattern = re.compile(
            r"^[ \t]*" + re.escape(field) + r"[\s_\-:]+([^\n]+)",
            re.IGNORECASE | re.MULTILINE,
        )
        m = pattern.search(text)
        if m:
            raw = m.group(1).strip().lstrip("_-: \t")
            # Skip if the "value" is obviously another field name
            if raw and not any(raw.lower() == f.lower() for f in _KNOWN_FIELD_NAMES):
                fields[field] = raw
    return fields


def _parse_field_then_value(text: str) -> dict[str, str]:
    """Match known field names that appear alone on one line, value on the next.

    Handles patterns like::

        Name
        Rahul Misra
        City
        Asansol
    """
    fields: dict[str, str] = {}
    lines = [ln.strip() for ln in text.splitlines()]
    for i, line in enumerate(lines):
        if i + 1 >= len(lines):
            break
        for field in _KNOWN_FIELD_NAMES:
            if line.lower() == field.lower() and field not in fields:
                next_line = lines[i + 1].strip().lstrip("_-: \t")
                if next_line and not any(
                    next_line.lower() == f.lower() for f in _KNOWN_FIELD_NAMES
                ):
                    fields[field] = next_line
                break
    return fields


def _parse_tab_separated(text: str) -> dict[str, str]:
    """Parse tab-separated ``FieldName\\tValue`` pairs."""
    fields: dict[str, str] = {}
    for line in text.splitlines():
        parts = line.split("\t", 1)
        if len(parts) == 2:
            fname = parts[0].strip()
            fval = parts[1].strip()
            if fname and fval and len(fname) <= 60:
                fields[fname] = fval
    return fields


def _parse_pdf(data: bytes) -> dict[str, str]:
    """Extract key-value pairs from PDF bytes using multiple strategies.

    Strategies tried in order (first non-empty result wins):

    1. PDF form widget fields (fillable form extraction via PyMuPDF)
    2. Address-book-aware parser (PDFService.map_address_book_fields)
    3. ``Field: Value`` or ``Field = Value`` lines (colon/equals separator)
    4. Known field names followed by their value on the same line
    5. Known field names on one line with the value on the next line
    6. Tab-separated ``FieldName\\tValue`` pairs
    7. Tesseract OCR text extraction (for scanned PDFs)
    8. Mindee / Koncile AI extraction (when API keys are configured)
    """
    import sys

    # Strategy 1: PDF form widget fields (most reliable for fillable forms)
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=data, filetype="pdf")
        widget_fields: dict[str, str] = {}
        for page in doc:
            for widget in page.widgets() or []:
                name = (widget.field_name or "").strip()
                value = widget.field_value or ""
                if isinstance(value, str):
                    value = value.strip()
                else:
                    value = str(value).strip()
                if name and value:
                    widget_fields[name] = value
        doc.close()
        if widget_fields:
            logger.debug("_parse_pdf: strategy 1 (widget fields) found %d fields", len(widget_fields))
            return widget_fields
    except Exception as exc:
        logger.debug("_parse_pdf: widget extraction failed: %s", exc)

    # Extract full text for text-based strategies
    full_text = _extract_pdf_text(data)
    full_text = full_text.strip() if full_text else ""

    # Strategy 2: Address-book-aware parser
    if full_text:
        try:
            _here = os.path.dirname(os.path.abspath(__file__))
            _backend = os.path.join(_here, "..", "backend")
            if _backend not in sys.path:
                sys.path.insert(0, os.path.abspath(_backend))
            from services.pdf_service import PDFService  # type: ignore[import]

            mapped = PDFService.map_address_book_fields(full_text)
            ab_fields = {item["field_name"]: item["value"] for item in mapped}
            if ab_fields:
                logger.debug("_parse_pdf: strategy 2 (address-book) found %d fields", len(ab_fields))
                return ab_fields
        except Exception as exc:
            logger.debug("_parse_pdf: address-book parser unavailable: %s", exc)

    if full_text:
        # Strategy 3: colon/equals separator
        result = _parse_txt(full_text)
        if result:
            logger.debug("_parse_pdf: strategy 3 (colon) found %d fields", len(result))
            return result

        # Strategy 4: known field names, value on same line
        result = _parse_known_fields_inline(full_text)
        if result:
            logger.debug("_parse_pdf: strategy 4 (inline keywords) found %d fields", len(result))
            return result

        # Strategy 5: known field names, value on next line
        result = _parse_field_then_value(full_text)
        if result:
            logger.debug("_parse_pdf: strategy 5 (next-line) found %d fields", len(result))
            return result

        # Strategy 6: tab-separated pairs
        result = _parse_tab_separated(full_text)
        if result:
            logger.debug("_parse_pdf: strategy 6 (tab-sep) found %d fields", len(result))
            return result

    # Strategy 7: Tesseract OCR for scanned PDFs (no embedded text layer)
    # Strategy 8: Mindee / Koncile advanced AI extraction (if API keys are configured)
    try:
        from backend.services.advanced_extraction_service import AdvancedExtractionService  # type: ignore[import]

        ocr_text = AdvancedExtractionService.extract_with_ocr(data)
        if ocr_text.strip():
            logger.debug("_parse_pdf: OCR extracted %d chars", len(ocr_text))

            result = _parse_txt(ocr_text)
            if result:
                logger.debug("_parse_pdf: OCR+colon found %d fields", len(result))
                return result

            result = _parse_known_fields_inline(ocr_text)
            if result:
                logger.debug("_parse_pdf: OCR+inline found %d fields", len(result))
                return result

            result = _parse_field_then_value(ocr_text)
            if result:
                logger.debug("_parse_pdf: OCR+next-line found %d fields", len(result))
                return result

            result = _parse_tab_separated(ocr_text)
            if result:
                logger.debug("_parse_pdf: OCR+tab-sep found %d fields", len(result))
                return result

            # LLM-based extraction on OCR text (if OpenAI API key is configured)
            openai_key = os.environ.get("OPENAI_API_KEY")
            if openai_key:
                result = AdvancedExtractionService.extract_with_llm(ocr_text, openai_key)
                if result:
                    logger.debug("_parse_pdf: OCR+LLM found %d fields", len(result))
                    return result

        # Mindee / Koncile advanced AI extraction
        mindee_key = os.environ.get("MINDEE_API_KEY")
        koncile_key = os.environ.get("KONCILE_API_KEY")
        if mindee_key or koncile_key:
            result = AdvancedExtractionService.extract_multi_strategy(
                data,
                mindee_key=mindee_key,
                koncile_key=koncile_key,
            )
            if result:
                logger.debug("_parse_pdf: strategy 8 (advanced AI) found %d fields", len(result))
                return result
    except Exception as exc:
        logger.debug("_parse_pdf: advanced extraction failed: %s", exc)

    logger.warning("_parse_pdf: all strategies exhausted, no fields found")
    return {}


def _parse_docx(data: bytes) -> dict[str, str]:
    """Extract key-value pairs from a .docx file."""
    try:
        import docx  # python-docx

        doc = docx.Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs]
        # Also include table cells
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                # Treat a two-column row as field:value
                if len(cells) == 2 and cells[0] and cells[1]:
                    lines.append(f"{cells[0]}: {cells[1]}")
                else:
                    lines.append("  ".join(cells))
        return _parse_txt("\n".join(lines))
    except Exception as exc:
        logger.warning("_parse_docx: failed to parse DOCX: %s", exc)
        return {}


def _parse_excel(data: bytes, ext: str) -> dict[str, str]:
    """Extract key-value pairs from an Excel file (.xlsx or .xls)."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        ws = wb.active
        fields: dict[str, str] = {}
        for row in ws.iter_rows(values_only=True):
            non_empty = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if len(non_empty) >= 2:
                # First cell = field name, second cell = value
                fname = non_empty[0]
                fval = non_empty[1]
                if fname and fval:
                    fields[fname] = fval
            elif len(non_empty) == 1:
                # Single cell: try "Field: Value" parsing
                m = re.match(r"^([^:=]+?)\s*[:=]\s*(.+)$", non_empty[0])
                if m:
                    fname = m.group(1).strip()
                    fval = m.group(2).strip()
                    if fname and fval:
                        fields[fname] = fval
        return fields
    except Exception as exc:
        logger.warning("_parse_excel: failed to parse Excel: %s", exc)
        return {}


# ---------------------------------------------------------------------------
# POST /api/v1/training/add
# ---------------------------------------------------------------------------

@training_bp.route("/api/v1/training/add", methods=["POST"])
@login_required
def add_training():
    """Save labeled training examples for a document.

    Accepts JSON body::

        {
          "document_id": 1,
          "fields": {
            "Name": "Rahul Misra",
            "City": "Asansol",
            ...
          }
        }

    Replaces any existing training examples for the document with the new
    set, then rebuilds RAG embeddings so that future extractions benefit from
    the training data.

    Returns::

        {"ok": true, "document_id": 1, "saved": [{"field_name": ..., "correct_value": ...}, ...]}
    """
    data = request.get_json(silent=True) or {}
    document_id = data.get("document_id")
    fields_param = data.get("fields")

    if not document_id:
        return jsonify({"ok": False, "error": "document_id is required"}), 400

    # Validate document exists before anything else
    doc = Document.query.get_or_404(document_id)

    # Resolve fields: use provided dict, auto-load from ExtractedField, or merge list
    if isinstance(fields_param, dict) and fields_param:
        fields: dict = fields_param
    else:
        # Auto-load from existing ExtractedField rows for this document
        ef_rows = ExtractedField.query.filter_by(document_id=document_id).all()
        fields = {ef.field_name: ef.value for ef in ef_rows}
        # If fields_param is a list, merge those in as overrides
        if isinstance(fields_param, list):
            for item in fields_param:
                if not isinstance(item, dict):
                    continue
                fname = str(item.get("field_name", "")).strip()
                fval = str(item.get("field_value", "")).strip()
                if fname and fval:
                    fields[fname] = fval
        if not fields:
            return jsonify({"ok": False, "error": "'fields' must be a non-empty object"}), 400

    # Replace existing examples for this document
    TrainingExample.query.filter_by(document_id=document_id).delete()

    saved: list[dict] = []
    for field_name, correct_value in fields.items():
        correct_value = str(correct_value).strip() if correct_value is not None else ""
        if not correct_value:
            continue  # skip blank values
        ex = TrainingExample(
            document_id=document_id,
            field_name=str(field_name).strip(),
            correct_value=correct_value,
        )
        db.session.add(ex)
        saved.append({"field_name": ex.field_name, "correct_value": ex.correct_value})

    _log(
        current_user.id,
        "add_training",
        "document",
        str(document_id),
        details=f"fields={list(fields.keys())}",
    )
    db.session.commit()

    # Rebuild RAG embeddings if RAGService is available
    _rebuild_rag_embeddings(document_id)

    return jsonify({"ok": True, "document_id": document_id, "added": len(saved), "saved": saved})


# ---------------------------------------------------------------------------
# GET /api/v1/training/examples
# ---------------------------------------------------------------------------

@training_bp.route("/api/v1/training/examples", methods=["GET"])
@login_required
def list_examples_json():
    """Return all training examples as JSON."""
    examples = TrainingExample.query.order_by(TrainingExample.created_at.desc()).all()
    return jsonify(
        {
            "count": len(examples),
            "examples": [ex.to_dict() for ex in examples],
        }
    )


# ---------------------------------------------------------------------------
# GET /api/v1/training/list
# ---------------------------------------------------------------------------

@training_bp.route("/api/v1/training/list", methods=["GET"])
@login_required
def list_training_json():
    """Return all training examples as JSON (alias for /api/v1/training/examples)."""
    examples = TrainingExample.query.order_by(TrainingExample.created_at.desc()).all()
    return jsonify(
        {
            "ok": True,
            "count": len(examples),
            "examples": [ex.to_dict() for ex in examples],
        }
    )


# ---------------------------------------------------------------------------
# POST /api/v1/training/apply/<doc_id>
# ---------------------------------------------------------------------------

@training_bp.route("/api/v1/training/apply/<int:doc_id>", methods=["POST"])
@login_required
def apply_training_to_document(doc_id: int):
    """Apply training data to a document by overwriting all its extracted fields.

    For each field in the document's ``extracted_fields`` table, this endpoint
    looks up the most common ``correct_value`` across all TrainingExample rows
    that share the same ``field_name``.  If a training value is found, the
    field's value is **overwritten** (not just filled in when blank).

    Returns::

        {
          "ok": true,
          "doc_id": 1,
          "updated": 5,
          "skipped": 2,
          "fields": [
            {"field_name": "Name", "new_value": "Rahul Misra", "updated": true},
            ...
          ]
        }

    HTTP 404 is returned when the document does not exist.
    HTTP 400 is returned when there are no training examples in the database.
    """
    doc = Document.query.get_or_404(doc_id)

    # Load all training examples from the database
    all_examples = TrainingExample.query.all()
    if not all_examples:
        return jsonify({"ok": False, "error": "No training examples found in the database."}), 400

    # Build a mapping: field_name → list of non-blank correct_value strings
    from collections import Counter

    training_by_field: dict[str, list[str]] = {}
    for ex in all_examples:
        fname = (ex.field_name or "").strip()
        fval = (ex.correct_value or "").strip()
        if fname and fval:
            training_by_field.setdefault(fname, []).append(fval)

    if not training_by_field:
        return jsonify({"ok": False, "error": "No non-blank training values found."}), 400

    # Apply the most common training value to each extracted field
    fields = ExtractedField.query.filter_by(document_id=doc_id).all()

    result_fields: list[dict] = []
    updated_count = 0
    skipped_count = 0

    for field in fields:
        fname = (field.field_name or "").strip()
        candidates = training_by_field.get(fname)
        if not candidates:
            skipped_count += 1
            result_fields.append({"field_name": fname, "new_value": field.value, "updated": False})
            continue

        # Pick the most common value among all training examples for this field
        most_common_val = Counter(candidates).most_common(1)[0][0]

        # Preserve original value before overwriting
        if not field.is_edited:
            field.original_value = field.value
        field.value = most_common_val
        field.confidence = 0.90
        field.is_edited = True

        updated_count += 1
        result_fields.append({"field_name": fname, "new_value": most_common_val, "updated": True})

    doc.status = "edited"
    _log(
        current_user.id,
        "apply_training",
        "document",
        str(doc_id),
        details=f"updated={updated_count} skipped={skipped_count}",
    )
    db.session.commit()

    return jsonify({
        "ok": True,
        "doc_id": doc_id,
        "updated": updated_count,
        "skipped": skipped_count,
        "fields": result_fields,
    })


# ---------------------------------------------------------------------------
# DELETE /api/v1/training/<id>
# ---------------------------------------------------------------------------

@training_bp.route("/api/v1/training/<int:example_id>", methods=["DELETE"])
@login_required
def delete_training_example(example_id: int):
    """Delete a single TrainingExample by its ID."""
    ex = TrainingExample.query.get_or_404(example_id)
    _log(
        current_user.id,
        "delete_training_example",
        "training_example",
        str(example_id),
        details=f"field_name={ex.field_name!r}",
    )
    db.session.delete(ex)
    db.session.commit()
    return jsonify({"ok": True, "deleted_id": example_id})


# ---------------------------------------------------------------------------
# POST /api/v1/training/save-roi
# ---------------------------------------------------------------------------

@training_bp.route("/api/v1/training/save-roi", methods=["POST"])
@login_required
def save_roi_training():
    """Save multiple ROI-annotated training examples for a document in one request.

    Accepts JSON body::

        {
          "document_id": 123,
          "page_number": 1,
          "examples": [
            {
              "field_name": "Name",
              "correct_value": "Rahul Misra",
              "x0": 0.1, "y0": 0.2, "x1": 0.5, "y1": 0.25,
              "engine": "pytesseract"
            },
            {
              "field_name": "Cell Phone",
              "correct_value": "7699888010",
              "x0": 0.6, "y0": 0.4, "x1": 0.9, "y1": 0.45
            }
          ]
        }

    For each example, any existing ``TrainingExample`` row for the same
    ``(document_id, page_number, field_name)`` is replaced (upsert semantics).
    All bounding-box coordinates must be in the range [0, 1].

    Returns::

        {"ok": true, "document_id": 123, "saved": 2}
    """
    data = request.get_json(silent=True) or {}
    document_id = data.get("document_id")
    page_number = int(data.get("page_number") or 1)
    examples_param = data.get("examples")

    if not document_id:
        return jsonify({"ok": False, "error": "document_id is required"}), 400
    if not isinstance(examples_param, list) or not examples_param:
        return jsonify({"ok": False, "error": "'examples' must be a non-empty list"}), 400

    # Validate document exists
    Document.query.get_or_404(document_id)

    saved_count = 0
    errors: list[str] = []

    for idx, item in enumerate(examples_param):
        if not isinstance(item, dict):
            errors.append(f"examples[{idx}]: must be an object")
            continue

        field_name = str(item.get("field_name", "")).strip()
        correct_value = str(item.get("correct_value", "")).strip()

        if not field_name:
            errors.append(f"examples[{idx}]: 'field_name' is required")
            continue
        if not correct_value:
            errors.append(f"examples[{idx}]: 'correct_value' is required")
            continue

        # Validate and coerce bounding-box values
        roi: dict[str, float | None] = {}
        for coord in ("x0", "y0", "x1", "y1"):
            raw = item.get(coord)
            if raw is None:
                roi[coord] = None
            else:
                try:
                    val = float(raw)
                except (TypeError, ValueError):
                    errors.append(f"examples[{idx}]: '{coord}' must be a number")
                    roi[coord] = None
                    continue
                if not (0.0 <= val <= 1.0):
                    errors.append(f"examples[{idx}]: '{coord}' must be between 0 and 1")
                    roi[coord] = None
                else:
                    roi[coord] = val

        engine = str(item.get("engine", "")).strip() or None
        anchor_text = str(item.get("anchor_text", "")).strip() or None

        # Delete existing row for same document + page + field to avoid duplicates
        TrainingExample.query.filter_by(
            document_id=document_id,
            page_number=page_number,
            field_name=field_name,
        ).delete()

        ex = TrainingExample(
            document_id=document_id,
            field_name=field_name,
            correct_value=correct_value,
            page_number=page_number,
            x0=roi.get("x0"),
            y0=roi.get("y0"),
            x1=roi.get("x1"),
            y1=roi.get("y1"),
            engine=engine,
            anchor_text=anchor_text,
        )
        db.session.add(ex)
        saved_count += 1

    if errors and saved_count == 0:
        return jsonify({"ok": False, "errors": errors}), 400

    _log(
        current_user.id,
        "save_roi_training",
        "document",
        str(document_id),
        details=f"page={page_number} saved={saved_count}",
    )
    db.session.commit()

    # Rebuild RAG embeddings if RAGService is available
    _rebuild_rag_embeddings(document_id)

    response: dict = {"ok": True, "document_id": document_id, "saved": saved_count}
    if errors:
        response["warnings"] = errors
    return jsonify(response)


@training_bp.route("/training/examples")
@login_required
def examples_list():
    """HTML view — display all training examples grouped by document."""
    examples = (
        TrainingExample.query
        .order_by(TrainingExample.document_id, TrainingExample.field_name)
        .all()
    )

    # Group by document
    grouped: dict[int, dict] = {}
    for ex in examples:
        doc_id = ex.document_id
        if doc_id not in grouped:
            doc = Document.query.get(doc_id)
            grouped[doc_id] = {
                "document": doc,
                "examples": [],
            }
        grouped[doc_id]["examples"].append(ex)

    return render_template(
        "training/examples_list.html",
        grouped=grouped,
        total=len(examples),
    )


# ---------------------------------------------------------------------------
# GET /training/upload-sample   — render the upload form
# POST /training/upload-sample  — process the form submission
# ---------------------------------------------------------------------------

@training_bp.route("/training/upload-sample", methods=["GET", "POST"])
@login_required
def upload_sample():
    """Form-based UI for uploading a named training sample with field-value pairs."""
    if request.method == "POST":
        sample_name = request.form.get("sample_name", "").strip()
        document_type = request.form.get("document_type", "").strip()
        upload_mode = request.form.get("upload_mode", "manual")

        if not sample_name:
            flash("Sample name is required.", "danger")
            return render_template(
                "training/upload_sample.html",
                document_types=DOCUMENT_TYPES,
            )

        fields: dict[str, str] = {}

        if upload_mode == "file":
            # --- File upload mode ---
            uploaded_file = request.files.get("sample_file")
            if not uploaded_file or not uploaded_file.filename:
                flash("Please select a file to upload.", "danger")
                return render_template(
                    "training/upload_sample.html",
                    document_types=DOCUMENT_TYPES,
                )
            ext = os.path.splitext(uploaded_file.filename)[1].lower()
            if ext not in ALLOWED_EXTENSIONS:
                flash(
                    f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
                    "danger",
                )
                return render_template(
                    "training/upload_sample.html",
                    document_types=DOCUMENT_TYPES,
                )
            fields = _extract_fields_from_file(uploaded_file)
            if not fields:
                flash(
                    "No field-value pairs could be extracted from the uploaded file. "
                    "Please check the file format or use manual entry.",
                    "warning",
                )
                return render_template(
                    "training/upload_sample.html",
                    document_types=DOCUMENT_TYPES,
                )

            # For file uploads: redirect to review page so the user can
            # validate and correct extracted fields before saving.
            flask_session["training_review"] = {
                "sample_name": sample_name,
                "document_type": document_type,
                "fields": fields,
            }
            return redirect(url_for("training.review_sample"))

        else:
            # --- Manual entry mode ---
            field_names = request.form.getlist("field_name[]")
            field_values = request.form.getlist("field_value[]")
            for fname, fval in zip(field_names, field_values):
                fname = fname.strip()
                fval = fval.strip()
                if fname and fval:
                    fields[fname] = fval

            if not fields:
                flash("Please provide at least one field name and value.", "danger")
                return render_template(
                    "training/upload_sample.html",
                    document_types=DOCUMENT_TYPES,
                )

        # --- Save training sample (manual entry path) ---
        return _save_training_sample(sample_name, document_type, upload_mode, fields)

    return render_template(
        "training/upload_sample.html",
        document_types=DOCUMENT_TYPES,
    )


# ---------------------------------------------------------------------------
# GET  /training/upload-sample/review — show extracted fields for validation
# POST /training/upload-sample/review — confirm and save after review
# ---------------------------------------------------------------------------

@training_bp.route("/training/upload-sample/review", methods=["GET", "POST"])
@login_required
def review_sample():
    """Show extracted fields after file upload for user validation before saving."""
    review_data = flask_session.get("training_review")
    if not review_data:
        flash("No pending sample to review. Please upload a file first.", "warning")
        return redirect(url_for("training.upload_sample"))

    if request.method == "POST":
        sample_name = review_data.get("sample_name", "")
        document_type = review_data.get("document_type", "")

        # Collect edited fields from the review form
        field_names = request.form.getlist("field_name[]")
        field_values = request.form.getlist("field_value[]")
        confirmed_fields: dict[str, str] = {}
        for fname, fval in zip(field_names, field_values):
            fname = fname.strip()
            fval = fval.strip()
            if fname and fval:
                confirmed_fields[fname] = fval

        # Also handle newly added fields from the review form
        new_names = request.form.getlist("new_field_name[]")
        new_values = request.form.getlist("new_field_value[]")
        for fname, fval in zip(new_names, new_values):
            fname = fname.strip()
            fval = fval.strip()
            if fname and fval:
                confirmed_fields[fname] = fval

        if not confirmed_fields:
            flash("Please keep at least one field.", "danger")
            return render_template(
                "training/review_fields.html",
                sample_name=sample_name,
                document_type=document_type,
                fields=review_data.get("fields", {}),
            )

        # Clear the session review data
        flask_session.pop("training_review", None)

        return _save_training_sample(sample_name, document_type, "file", confirmed_fields)

    return render_template(
        "training/review_fields.html",
        sample_name=review_data.get("sample_name", ""),
        document_type=review_data.get("document_type", ""),
        fields=review_data.get("fields", {}),
    )
# POST /training/extract-preview — extract fields from an uploaded file (JSON)
# ---------------------------------------------------------------------------

@training_bp.route("/training/extract-preview", methods=["POST"])
@login_required
def extract_preview():
    """Extract field-value pairs from an uploaded file and return them as JSON.

    This endpoint is called via AJAX to let the user preview and edit the
    extracted fields *before* the training sample is saved.

    Returns::

        {"ok": true, "fields": {"Name": "Rahul Misra", ...}, "count": 6}

    or::

        {"ok": false, "error": "..."}
    """
    uploaded_file = request.files.get("sample_file")
    if not uploaded_file or not uploaded_file.filename:
        return jsonify({"ok": False, "error": "No file provided."}), 400

    ext = os.path.splitext(uploaded_file.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({
            "ok": False,
            "error": f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        }), 400

    fields = _extract_fields_from_file(uploaded_file)

    if not fields:
        return jsonify({
            "ok": False,
            "error": (
                "No field-value pairs could be extracted from the uploaded file. "
                "The file may be scanned/image-only, empty, or in an unsupported layout. "
                "Please use Manual Entry instead."
            ),
        }), 422

    return jsonify({"ok": True, "fields": fields, "count": len(fields)})


# ---------------------------------------------------------------------------
# POST /training/examples/<doc_id>/delete — remove a sample and its examples
# ---------------------------------------------------------------------------

@training_bp.route("/training/examples/<int:doc_id>/delete", methods=["POST"])
@login_required
def delete_sample(doc_id: int):
    """Delete all training examples for a document (and the document if it is a training placeholder)."""
    doc = Document.query.get_or_404(doc_id)

    # Remove all training examples linked to this document
    deleted_count = TrainingExample.query.filter_by(document_id=doc_id).delete()

    _log(
        current_user.id,
        "delete_sample",
        "document",
        str(doc_id),
        details=f"deleted {deleted_count} training example(s) for doc '{doc.filename}'",
    )

    # Remove the document record if it was a training placeholder
    if doc.status == "training":
        db.session.delete(doc)

    db.session.commit()

    flash(f"Sample deleted ({deleted_count} field(s) removed).", "success")
    return redirect(url_for("training.examples_list"))


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _save_training_sample(
    sample_name: str,
    document_type: str,
    upload_mode: str,
    fields: dict[str, str],
):
    """Persist a training sample to the database and redirect to the examples list.

    Creates a Document placeholder (status='training') and TrainingExample rows
    for every field-value pair, then flashes a success message.
    """
    safe_name = secure_filename(sample_name) or "sample"
    placeholder_path = os.path.join(
        current_app.config.get("UPLOAD_FOLDER", "uploads"),
        f"{safe_name}.training",
    )
    display_name = f"{sample_name} [{document_type}]" if document_type else sample_name
    doc = Document(
        filename=display_name,
        file_path=placeholder_path,
        status="training",
        uploaded_by=current_user.id,
    )
    db.session.add(doc)
    db.session.flush()  # get doc.id before committing

    saved: list[dict] = []
    for fname, fval in fields.items():
        ex = TrainingExample(
            document_id=doc.id,
            field_name=fname,
            correct_value=fval,
            created_by=current_user.id,
        )
        db.session.add(ex)
        saved.append({"field_name": fname, "correct_value": fval})

    _log(
        current_user.id,
        "upload_sample",
        "document",
        str(doc.id),
        details=f"sample_name={sample_name!r}, document_type={document_type!r}, mode={upload_mode!r}, fields={list(fields.keys())}",
    )
    db.session.commit()

    flash(
        f"Sample '{sample_name}' saved with {len(saved)} field(s).",
        "success",
    )
    return redirect(url_for("training.examples_list"))


def _rebuild_rag_embeddings(document_id: int) -> None:
    """Attempt to rebuild RAG embeddings for *document_id* after training data update."""
    try:
        import os
        from services.rag_service import RAGService  # type: ignore[import]

        doc = Document.query.get(document_id)
        if doc is None or not doc.file_path:
            return

        import sys
        _root = current_app.root_path
        _backend = os.path.join(_root, "backend")
        if _backend not in sys.path:
            sys.path.append(_backend)

        from services.pdf_service import PDFService  # type: ignore[import]

        pdf_svc = PDFService()
        if not os.path.exists(doc.file_path):
            return

        text, _tables, _pages = pdf_svc.extract(doc.file_path)
        rag_dir = os.path.join(_root, os.environ.get("RAG_DIR", "rag_data"))
        rag_svc = RAGService(rag_dir=rag_dir)
        rag_svc.save_rag_text(str(document_id), text)
        current_app.logger.info(
            "TrainingService: rebuilt RAG embeddings for doc %s", document_id
        )
    except Exception as exc:  # pragma: no cover
        current_app.logger.warning(
            "TrainingService: could not rebuild RAG embeddings for doc %s: %s",
            document_id,
            exc,
        )


def _log(
    user_id: int,
    action: str,
    resource_type: str,
    resource_id: str,
    details: str = "",
) -> None:
    """Insert an AuditLog entry (caller must commit the session)."""
    entry = AuditLog(
        user_id=user_id,
        action=action,
        resource_type=resource_type,
        resource_id=resource_id,
        details=details or None,
    )
    db.session.add(entry)
